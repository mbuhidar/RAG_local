"""
Document Loader Module

Loads documents from various file formats including PDF, DOCX, TXT, and Markdown.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

from . import Document

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Loads documents from local file system."""

    def __init__(self, supported_formats: Optional[List[str]] = None):
        """
        Initialize document loader.

        Args:
            supported_formats: List of supported file extensions
        """
        self.supported_formats = supported_formats or ['txt', 'pdf', 'md', 'docx']

    def load_file(self, file_path: Path) -> Optional[Document]:
        """
        Load a single file.

        Args:
            file_path: Path to the file

        Returns:
            Document object or None if failed
        """
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None

        extension = file_path.suffix.lower().lstrip('.')
        if extension not in self.supported_formats:
            logger.warning(f"Unsupported format: {extension}")
            return None

        try:
            if extension == 'txt' or extension == 'md':
                return self._load_text_file(file_path)
            elif extension == 'pdf':
                return self._load_pdf_file(file_path)
            elif extension == 'docx':
                return self._load_docx_file(file_path)
        except Exception as e:
            logger.error(f"Error loading file {file_path}: {e}")
            return None

    def load_directory(
        self,
        directory_path: Path,
        recursive: bool = True
    ) -> List[Document]:
        """
        Load all supported documents from a directory.

        Args:
            directory_path: Path to directory
            recursive: Whether to search subdirectories

        Returns:
            List of Document objects
        """
        documents = []
        pattern = "**/*" if recursive else "*"

        for file_path in directory_path.glob(pattern):
            if file_path.is_file():
                doc = self.load_file(file_path)
                if doc:
                    documents.append(doc)

        logger.info(f"Loaded {len(documents)} documents from {directory_path}")
        return documents

    def _load_text_file(self, file_path: Path) -> Document:
        """Load plain text or markdown file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        metadata = self._extract_file_metadata(file_path)
        return Document(content=content, metadata=metadata)

    def _load_pdf_file(self, file_path: Path) -> Document:
        """Load PDF file."""
        try:
            import pypdf
        except ImportError:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            raise

        content_parts = []
        with open(file_path, 'rb') as f:
            pdf_reader = pypdf.PdfReader(f)
            num_pages = len(pdf_reader.pages)

            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                content_parts.append(text)

        content = "\n\n".join(content_parts)
        metadata = self._extract_file_metadata(file_path)
        metadata['num_pages'] = num_pages
        metadata['format'] = 'pdf'

        return Document(content=content, metadata=metadata)

    def _load_docx_file(self, file_path: Path) -> Document:
        """Load DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise

        doc = DocxDocument(file_path)
        content_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        content = "\n\n".join(content_parts)

        metadata = self._extract_file_metadata(file_path)
        metadata['num_paragraphs'] = len(doc.paragraphs)
        metadata['format'] = 'docx'

        return Document(content=content, metadata=metadata)

    def _extract_file_metadata(self, file_path: Path) -> Dict:
        """Extract metadata from file."""
        stat = file_path.stat()
        return {
            'filename': file_path.name,
            'filepath': str(file_path.absolute()),
            'file_size': stat.st_size,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'extension': file_path.suffix.lower().lstrip('.')
        }
